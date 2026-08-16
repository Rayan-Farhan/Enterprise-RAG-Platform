"""Unit tests for FormatRouter (Task 1.5)."""

from pathlib import Path
from unittest.mock import MagicMock

import docx
import pytest

from app.ingestion.parsers.base import (
    DocumentParser,
    ElementType,
    ParsedDocument,
    ParsedElement,
    ParsedPage,
)
from app.ingestion.parsers.router import FormatRouter

CORPUS_DIR = Path("benchmarks/corpus")


@pytest.fixture
def sample_pdf() -> Path:
    return CORPUS_DIR / "organizational_structure.pdf"


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    p = tmp_path / "sample.docx"
    doc = docx.Document()
    doc.add_heading("Sample Policy", level=1)
    doc.add_paragraph("Policy content text.")
    doc.save(str(p))
    return p


def test_format_router_routes_pdf(sample_pdf: Path) -> None:
    if not sample_pdf.exists():
        pytest.skip("Corpus file not generated.")

    router = FormatRouter()
    doc = router.route_and_parse(sample_pdf)
    assert doc.parser_name == "docling"
    assert doc.total_pages == 1


def test_format_router_routes_docx(sample_docx: Path) -> None:
    router = FormatRouter()
    doc = router.route_and_parse(sample_docx)
    assert doc.parser_name == "office_parser"
    assert doc.file_type == "docx"


def test_format_router_fallback_on_error(sample_pdf: Path) -> None:
    if not sample_pdf.exists():
        pytest.skip("Corpus file not generated.")

    failing_parser = MagicMock(spec=DocumentParser)
    failing_parser.parser_name = "failing_mock"
    failing_parser.parse.side_effect = ValueError("Corrupt PDF stream")

    fallback_parser = MagicMock(spec=DocumentParser)
    fallback_parser.parser_name = "fallback_mock"
    fallback_doc = ParsedDocument(
        filename="test.pdf",
        file_type="pdf",
        total_pages=1,
        pages=[
            ParsedPage(
                page_number=1,
                elements=[
                    ParsedElement(
                        element_id="el_1",
                        element_type=ElementType.PARAGRAPH,
                        text="Recovered text",
                        page_number=1,
                    )
                ],
            )
        ],
        parser_name="fallback_mock",
    )
    fallback_parser.parse.return_value = fallback_doc

    router = FormatRouter(
        pdf_primary=failing_parser,
        pdf_fallbacks=[fallback_parser],
    )

    doc = router.route_and_parse(sample_pdf)
    assert doc.parser_name == "fallback_mock"
    assert "prior_parser_errors" in doc.metadata
    assert "Corrupt PDF stream" in doc.metadata["prior_parser_errors"][0]


def test_format_router_file_not_found() -> None:
    router = FormatRouter()
    with pytest.raises(FileNotFoundError):
        router.route_and_parse("non_existent_file.pdf")
